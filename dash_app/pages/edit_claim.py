# pages/edit_claim.py
import dash
from dash import html, dcc, Input, Output, State, callback
import dash_mantine_components as dmc
import pymysql
import requests  # Imported for fetching the DOCX file
from db import get_db_connection
import time
from io import BytesIO
import re
from docx import Document
from dotenv import load_dotenv
import json
from dash_iconify import DashIconify
import boto3
from botocore.config import Config
import os
from urllib.parse import urlparse

# Load environment variables from .env file
load_dotenv('/opt/PrelimSite/.env')

dash.register_page(__name__, path_template="/edit/<cid>")

def layout(cid=None, **other_kwargs):
    """
    This layout function receives `cid` from the URL, queries the DB for that record,
    and populates the form fields with existing data.
    """

    # Initialize empty dict; if CID is valid, we'll overwrite with DB data
    claim_data = {}

    # Convert cid to int if needed, then query the DB
    if cid is not None:
        try:
            claim_id = int(cid)
        except ValueError:
            claim_id = None

        if claim_id:
            conn = None
            cursor = None
            try:
                conn = get_db_connection()
                cursor = conn.cursor(pymysql.cursors.DictCursor)
                sql = "SELECT * FROM claims WHERE id = %s"
                cursor.execute(sql, (claim_id,))
                row = cursor.fetchone()
                if row:
                    claim_data = row
                    print("DEBUG: Coverage Values from DB:")
                    print(f"Coverage A Deductible: {row.get('Coverage_A_Deductible')}")
                    print(f"Coverage A Reserve: {row.get('Coverage_A_Reserve')}")
                    print(f"Coverage A Advance: {row.get('Coverage_A_Advance')}")
                    print(f"Coverage B Deductible: {row.get('Coverage_B_Deductible')}")
                    print(f"Coverage B Reserve: {row.get('Coverage_B_Reserve')}")
                    print(f"Coverage B Advance: {row.get('Coverage_B_Advance')}")
            except Exception as e:
                print(f"Error fetching claim data for CID={cid}: {e}")
            finally:
                if cursor:
                    cursor.close()
                if conn:
                    conn.close()

    # We store cid in a dcc.Store so we can access it in the callback
    store_cid = dcc.Store(id="cid-store", data=cid)

    # Parse JSON output from the correct column name
    try:
        json_data = json.loads(claim_data['confidence_json']) if claim_data.get('confidence_json') else {
            "entities": []
        }
    except (json.JSONDecodeError, TypeError):
        print(f"Warning: Invalid confidence_json data: {claim_data.get('confidence_json')}")
        json_data = {
            "entities": []
        }

    # Create a mapping between your form fields and JSON types
    field_type_mapping = {
        "coverage-a-reserve": "Coverage-A-Building_Reserve",
        "coverage-a": "Coverage-A-Building",
        "coverage-a-deductible": "Coverage-A-Building_Deductible",
        "coverage-a-advance": "Coverage-A-Building_Advance",
        "coverage-b": "Coverage-B-Contents",
        "coverage-b-reserve": "Coverage-B-Contents_Reserve",
        "coverage-b-deductible": "Coverage-B-Contents_Deductible",
        "coverage-b-advance": "Coverage-B-Contents_Advance",
        "policyholder": "Policy_Holder",
        "claim-number": "File_Number",
        "date-of-loss": "Date_Of_Loss",
        "insurer": "Insurer",
        "adjuster-name": "Adjuster",
        "policy-number": "Policy_Number",
        "claim-type": "Claim_Type",
        "contact-info-adjuster": "Adjuster_Email",
        "contact-info-insured": "Policyholder_Email",
        "loss-address": "Property_Address",
        "claim-assigned-date": "Date_Assigned",
        "claim-contact-date": "Date_Contacted",
        "claim-inspection-date": "Date_Inspected",
        "dwelling-unit-rcv-loss": "DwellingUnit_Insured_Damage_RCV_Loss",
        "detached-garage-rcv-loss": "DetachedGarage_Insured_Damage_RCV_Loss",
        "improvements-rcv-loss": "Improvements_Insured_Damage_RCV_Loss",
        "contents-rcv-loss": "Contents_Insured_Damage_RCV_Loss",
        # Add more mappings as needed
    }

    def get_confidence(field_id):
        # Get the corresponding JSON type for this field
        json_type = field_type_mapping.get(field_id)
        print(f"\nDEBUG get_confidence:")
        print(f"Field ID: {field_id}")
        print(f"Mapped JSON type: {json_type}")
        print(f"json_data: {json_data}")
        
        if not json_type:
            print(f"No mapping found for field_id '{field_id}', returning default 0.90")
            return .90  # Default confidence if no mapping exists
        
        if json_data and 'entities' in json_data:
            print(f"Searching through {len(json_data['entities'])} entities")
            for entity in json_data['entities']:
                print(f"  Checking entity: type={entity.get('type')}, confidence={entity.get('confidence')}")
                if entity.get('type') == json_type:
                    confidence = entity.get('confidence', 1.0)
                    print(f"  ✓ Match found! Returning confidence: {confidence}")
                    return confidence
            print(f"No matching entity found for type '{json_type}', returning 1.0")
        else:
            print("No json_data or 'entities' not in json_data, returning 1.0")
            print(f"json_data type: {type(json_data)}")
            if json_data:
                print(f"json_data keys: {json_data.keys()}")
        return 1.0

    def get_style(field_id, base_style=None):
        base_style = base_style or {}
        
        # Map component IDs to database column names
        field_mapping = {
            "coverage-a": "coverage_building",
            "coverage-b": "coverage_contents",
            "coverage-a-deductible": "Coverage_A_Deductible",
            "coverage-b-deductible": "Coverage_B_Deductible",
            "coverage-a-reserve": "Coverage_A_Reserve",
            "coverage-b-reserve": "Coverage_B_Reserve",
            "coverage-a-advance": "Coverage_A_Advance",
            "coverage-b-advance": "Coverage_B_Advance",
            "policyholder": "Policyholder",
            "claim-number": "claim_number",
            "date-of-loss": "Date_Of_Loss",
            "insurer": "Insurer",
            "adjuster-name": "Adjuster_Name",
            "policy-number": "Policy_Number",
            "claim-type": "claim_type",
            "contact-info-adjuster": "Adjuster_Contact_Info",
            "contact-info-insured": "Insured_Contact_Info",
            "loss-address": "Loss_Address",
            "claim-assigned-date": "Claim_Assigned_Date",
            "claim-contact-date": "Claim_Contact_Date",
            "claim-inspection-date": "Claim_Inspection_Date",
            "dwelling-unit-rcv-loss": "DwellingUnit_Insured_Damage_RCV_Loss",
            "detached-garage-rcv-loss": "DetachedGarage_Insured_Damage_RCV_Loss",
            "improvements-rcv-loss": "Improvements_Insured_Damage_RCV_Loss",
            "contents-rcv-loss": "Contents_Insured_Damage_RCV_Loss"
        }
        
        # Common styles for all states
        base_style.update({
            "borderRadius": "8px",
            "borderWidth": "1px",
            "borderStyle": "solid",
            "transition": "all 0.2s ease-in-out",
            "boxShadow": "0 1px 3px rgba(0,0,0,0.1)",
            "padding": "8px 12px",  # Added padding inside the input
            "marginBottom": "10px",  # Space between inputs
            "input": {  # Styles for the actual input element
                "padding": "8px",
                "fontSize": "14px",
                "lineHeight": "1.5",
            }
        })
        
        # Get the correct database column name
        db_field = field_mapping.get(field_id, field_id)
        value = claim_data.get(db_field)
        
        if value is None:
            value = ""
        elif isinstance(value, str):
            value = value.strip()
        
        confidence = get_confidence(field_id)
        
        if not value:
            base_style.update({
                "backgroundColor": "#fff8e1",
                "borderColor": "#ffd54f",
                "&:hover": {
                    "borderColor": "#ffb300",
                    "boxShadow": "0 2px 4px rgba(255,179,0,0.15)"
                },
                "input": {
                    "padding": "8px",
                    "fontSize": "14px",
                    "lineHeight": "1.5",
                    "backgroundColor": "transparent"  # Ensure input background is transparent
                }
            })
        elif confidence < 0.98:
            base_style.update({
                "backgroundColor": "#ffebee",
                "borderColor": "#ef9a9a",
                "&:hover": {
                    "borderColor": "#ef5350",
                    "boxShadow": "0 2px 4px rgba(239,83,80,0.15)"
                },
                "input": {
                    "padding": "8px",
                    "fontSize": "14px",
                    "lineHeight": "1.5",
                    "backgroundColor": "transparent"
                }
            })
        else:
            base_style.update({
                "backgroundColor": "#f1f8e9",
                "borderColor": "#aed581",
                "&:hover": {
                    "borderColor": "#7cb342",
                    "boxShadow": "0 2px 4px rgba(124,179,66,0.15)"
                },
                "input": {
                    "padding": "8px",
                    "fontSize": "14px",
                    "lineHeight": "1.5",
                    "backgroundColor": "transparent"
                }
            })
            
        return base_style

    def get_text_input(label, id, value, placeholder, style):
        return dmc.TextInput(
            label=label,
            id=id,
            value=value,
            placeholder=placeholder,
            style=style
        )
    
    def get_textarea(label, id, value, placeholder):
        return dmc.Textarea(
            label=label,
            id=id,
            value=value,
            placeholder=placeholder,
            minRows=3,
        )
    return dmc.Stack(
        [
            store_cid,  # hidden, just holds the cid for the callback

            # Create a Group to hold the heading, button, and tooltip
            dmc.Group(
                [
                    # Left side with heading and tooltip
                    dmc.Group(
                        [
                            html.H3(f"Editing Claim: {cid}", style={"margin": 0}),
                            dmc.Tooltip(
                                label=dmc.List(
                                    spacing="xs",
                                    size="sm",
                                    children=[
                                        dmc.ListItem(
                                            "Green: Verified data (confidence > 96%). Information is most likely correct.",
                                            icon=dmc.ThemeIcon(
                                                DashIconify(icon="radix-icons:check", width=16),
                                                radius="xl",
                                                color="green",
                                                size=24,
                                            ),
                                        ),
                                        dmc.ListItem(
                                            "Yellow: Empty field needs attention. May or may not need to be filled out.",
                                            icon=dmc.ThemeIcon(
                                                DashIconify(icon="radix-icons:question-mark", width=16),
                                                radius="xl",
                                                color="yellow",
                                                size=24,
                                            ),
                                        ),
                                        dmc.ListItem(
                                            "Red: Low confidence data. Program not confident in this information. (< 96%)",
                                            icon=dmc.ThemeIcon(
                                                DashIconify(icon="radix-icons:exclamation-triangle", width=16),
                                                radius="xl",
                                                color="red",
                                                size=24,
                                            ),
                                        ),
                                    ],
                                ),
                                multiline=True,
                                maw=300,
                                children=dmc.ThemeIcon(
                                    DashIconify(icon="radix-icons:question-mark", width=16),
                                    radius="xl",
                                    color="gray",
                                    size=24,
                                    style={"cursor": "help"}
                                ),
                            ),
                        ],
                        gap="xs",
                    ),
                    # Right side button - now with _blank target
                    html.A(
                        dmc.Button(
                            "View Binder PDF",
                            color="blue",
                            variant="filled",
                            leftSection=html.I(className="fas fa-file-pdf"),
                        ),
                        id="view-binder-link",
                        href="#",
                        target="_blank",  # Changed back to _blank to open in new window
                        rel="noopener noreferrer",  # Added for security best practices
                        style={"textDecoration": "none"}
                    ),
                ],
                justify="space-between",
                align="center",
            ),
            
            dmc.Text("Please fill out the fields below (data loaded from DB)."),

            # ========== Basic Fields in columns ==========
            dmc.Group(
                [
                    get_text_input("Claim Number", "claim-number", claim_data.get("claim_number", ""), "Enter claim number", get_style("claim-number", {"width": "45%"}),),
                    get_text_input("Policyholder", "policyholder", claim_data.get("Policyholder", ""), "Enter policyholder name", get_style("policyholder", {"width": "45%"}),),
                ],
                justify="space-between",
                style={"marginTop": "1rem"}
            ),

            dmc.Group(
                [
                    get_text_input("Loss Address", "loss-address", claim_data.get("Loss_Address", ""), "Enter loss address", get_style("loss-address", {"width": "45%"}),),
                    get_text_input("Date of Loss", "date-of-loss", claim_data.get("Date_Of_Loss", ""), "YYYY-MM-DD", get_style("date-of-loss", {"width": "45%"}),),
                ],
                justify="space-between"
            ),

            dmc.Group(
                [
                    get_text_input("Insurer", "insurer", claim_data.get("Insurer", ""), "e.g. Acme Insurance", get_style("insurer", {"width": "45%"}),),
                    get_text_input("Adjuster Name", "adjuster-name", claim_data.get("Adjuster_Name", ""), "e.g. John Doe", get_style("adjuster-name", {"width": "45%"}),),
                ],
                justify="space-between"
            ),

            dmc.Group(
                [
                    get_text_input("Policy Number", "policy-number", claim_data.get("Policy_Number", ""), "Policy #", get_style("policy-number", {"width": "45%"}),),
                    get_text_input("Claim Type", "claim-type", claim_data.get("claim_type", ""), "Building Only/Contents Only/Building and Contents", get_style("claim-type", {"width": "45%"}),),
                ],
                justify="space-between"
            ),

            dmc.Group(
                [
                    get_text_input("Contact Information (Adjuster)", "contact-info-adjuster", claim_data.get("Adjuster_Contact_Info", ""), "Phone / Email", get_style("contact-info-adjuster", {"width": "45%"}),),
                    get_text_input("Contact Information (Insured)", "contact-info-insured", claim_data.get("Insured_Contact_Info", ""), "Phone / Email", get_style("contact-info-insured", {"width": "45%"}),),
                ],
                justify="space-between"
            ),

            # ========== Coverage Fields in three columns ==========
            dmc.Text("Coverage A - Building", size="lg", fw=500, style={"marginTop": "1rem"}),
            dmc.SimpleGrid(
                cols=3,
                spacing="xl",
                children=[
                    # Column 1: Coverage and Deductible
                    dmc.Stack(
                        [
                            get_text_input("Coverage A", "coverage-a", claim_data.get("coverage_building", ""), "Building Coverage", get_style("coverage-a")),
                            get_text_input("Coverage A Deductible", "coverage-a-deductible", claim_data.get("Coverage_A_Deductible", ""), "", get_style("coverage-a-deductible")),
                        ],
                    ),
                    # Column 2: RCV Loss Values
                    dmc.Stack(
                        [
                            get_text_input("Dwelling Unit RCV Loss", "dwelling-unit-rcv-loss", claim_data.get("DwellingUnit_Insured_Damage_RCV_Loss", ""), "Enter RCV Loss amount", get_style("dwelling-unit-rcv-loss")),
                            get_text_input("Detached Garage RCV Loss", "detached-garage-rcv-loss", claim_data.get("DetachedGarage_Insured_Damage_RCV_Loss", ""), "Enter RCV Loss amount", get_style("detached-garage-rcv-loss")),
                        ],
                    ),
                    # Column 3: Reserve and Advance
                    dmc.Stack(
                        [
                            get_text_input("Coverage A Reserve", "coverage-a-reserve", claim_data.get("Coverage_A_Reserve", ""), "", get_style("coverage-a-reserve")),
                            get_text_input("Coverage A Advance", "coverage-a-advance", claim_data.get("Coverage_A_Advance", ""), "", get_style("coverage-a-advance")),
                        ],
                    ),
                ],
            ),

            # Coverage B Section
            dmc.Text("Coverage B - Contents", size="lg", fw=500, style={"marginTop": "2rem"}),
            dmc.SimpleGrid(
                cols=3,
                spacing="xl",
                children=[
                    # Column 1: Coverage and Deductible
                    dmc.Stack(
                        [
                            get_text_input("Coverage B", "coverage-b", claim_data.get("coverage_contents", ""), "Contents Coverage", get_style("coverage-b")),
                            get_text_input("Coverage B Deductible", "coverage-b-deductible", claim_data.get("Coverage_B_Deductible", ""), "", get_style("coverage-b-deductible")),
                        ],
                    ),
                    # Column 2: RCV Loss Values
                    dmc.Stack(
                        [
                            get_text_input("Contents RCV Loss", "contents-rcv-loss", claim_data.get("Contents_Insured_Damage_RCV_Loss", ""), "Enter RCV Loss amount", get_style("contents-rcv-loss")),
                            get_text_input("Improvements RCV Loss", "improvements-rcv-loss", claim_data.get("Improvements_Insured_Damage_RCV_Loss", ""), "Enter RCV Loss amount", get_style("improvements-rcv-loss")),
                        ],
                    ),
                    # Column 3: Reserve and Advance
                    dmc.Stack(
                        [
                            get_text_input("Coverage B Reserve", "coverage-b-reserve", claim_data.get("Coverage_B_Reserve", ""), "", get_style("coverage-b-reserve")),
                            get_text_input("Coverage B Advance", "coverage-b-advance", claim_data.get("Coverage_B_Advance", ""), "", get_style("coverage-b-advance")),
                        ],
                    ),
                ],
            ),

            # ========== Paragraph Field (full width) ==========
            get_textarea("Current Claim Status Paragraph", "Current_Claim_Status_Par", claim_data.get("Current_Claim_Status_Par", ""), "Summarize the current claim status..."),

            # ========== NEW FIELDS: Dates in columns ==========
            # Here we have 3 date fields, let's place them in a single row
            dmc.Group(
                [
                    get_text_input("Claim Assigned Date", "claim-assigned-date", claim_data.get("Claim_Assigned_Date", ""), "YYYY-MM-DD", get_style("claim-assigned-date", {"width": "30%"})),
                    get_text_input("Claim Contact Date", "claim-contact-date", claim_data.get("Claim_Contact_Date", ""), "YYYY-MM-DD", get_style("claim-contact-date", {"width": "30%"})),
                    get_text_input("Claim Inspection Date", "claim-inspection-date", claim_data.get("Claim_Inspection_Date", ""), "YYYY-MM-DD", get_style("claim-inspection-date", {"width": "30%"}))
                ],
                justify="space-between",
                style={"marginTop": "1rem"}
            ),

            # ========== Paragraph-type fields (full width) ==========
            get_textarea("Preliminary Report Paragraph", "preliminary-report-par", claim_data.get("Preliminary_Report_Par", ""), "Details for Preliminary Report..."),

            get_textarea("Insured Communication Paragraph", "insured-communication-paragraph", claim_data.get("Insured_Communication_Paragraph", ""), "Details about communication with the insured..."),
                        
            get_textarea("Claim Reserve Paragraph", "claim-reserve-paragraph", claim_data.get("Claim_Reserve_Paragraph", ""), "Details about the claim reserves..."),

            get_textarea("Insured Concern Paragraph", "insured-concern-paragraph", claim_data.get("Insured_Concern_Paragraph", ""), "Summarize any insured concerns..."),

            get_textarea("Adjuster Response Paragraph", "adjuster-response-paragraph", claim_data.get("Adjuster_Response_Paragraph", ""), "Adjuster's response or actions taken..."),

            get_textarea("Supporting Documents Paragraph", "supporting-doc-paragraph", claim_data.get("Supporting_Doc_Paragraph", ""), "Summary of supporting documents..."),

            get_textarea("Next Steps Paragraph", "next-steps-paragraph", claim_data.get("Next_Steps_Paragraph", ""), "Outline the next steps in the claim process..."),
            
            get_textarea("Final Report Paragraph", "final-report-paragraph", claim_data.get("Final_Report_Paragraph", ""), "Details of the final report..."),

            get_textarea("Claim Summary Paragraph", "claim-summary-par", claim_data.get("Claim_Summary_Par", ""), "A concise summary of the claim..."),

            dmc.Select(
            label="Review Status",
            id="review-status",
            value=claim_data.get("Review_Status", "In Review"),
            data=[
                {"label": "In Review",  "value": "In Review"},
                {"label": "Reviewed",   "value": "Reviewed"},
                
                {"label": "Needs Work", "value": "Needs Work"},
            ],
            style={"width": "45%"}
        ),

            # ========== Save Button & Confirmation ==========
            dmc.Group(
                [
                    dmc.Button("Save Changes", id="save-button", color="blue", variant="filled"),
                ],
                justify="flex-end",
                style={"marginTop": "1rem"}
            ),
            html.Div(
                id="save-confirmation",
                style={"color": "green", "marginTop": "1rem"}
            ),
            

            # ======= Download Report Button =======
            dmc.Button(
                "Download Report",
                id="download-docx-button",
                n_clicks=0,
                color="blue",
                variant="filled",
                fullWidth=True,
                mt="md",
            ),

            # ======= Download Component =======
            dcc.Download(id="download-docx"),

            # ======= Notification Container =======
            dmc.Notification(
                id="download-notification",
                title="Download Complete",
                message="Your file has been successfully downloaded.",
                color="green",
                autoClose=5000,
                action={
                    "label": "Close",
                    "onClick": "function() { return false; }"
                }
            ),

        ],  # Close the list
        style={"padding": "20px"}
    )  # Close the Stack

@callback(
    Output("save-confirmation", "children"),
    Input("save-button", "n_clicks"),
    [
        State("cid-store", "data"),
        # Existing Fields
        State("claim-number", "value"),
        State("policyholder", "value"),
        State("loss-address", "value"),
        State("date-of-loss", "value"),
        State("insurer", "value"),
        State("adjuster-name", "value"),
        State("policy-number", "value"),
        State("claim-type", "value"),
        State("contact-info-adjuster", "value"),
        State("contact-info-insured", "value"),
        State("coverage-a", "value"),
        State("coverage-a-deductible", "value"),
        State("coverage-a-reserve", "value"),
        State("coverage-a-advance", "value"),
        State("coverage-b", "value"),
        State("coverage-b-deductible", "value"),
        State("coverage-b-reserve", "value"),
        State("coverage-b-advance", "value"),
        State("Current_Claim_Status_Par", "value"),
        # New Fields
        State("claim-assigned-date", "value"),
        State("claim-contact-date", "value"),
        State("claim-inspection-date", "value"),
        State("preliminary-report-par", "value"),
        State("insured-communication-paragraph", "value"),
        State("claim-reserve-paragraph", "value"),
        State("insured-concern-paragraph", "value"),
        State("adjuster-response-paragraph", "value"),
        State("supporting-doc-paragraph", "value"),
        State("next-steps-paragraph", "value"),
        State("final-report-paragraph", "value"),
        State("claim-summary-par", "value"),
        State("review-status", "value"),
        # Add new RCV Loss states
        State("dwelling-unit-rcv-loss", "value"),
        State("detached-garage-rcv-loss", "value"),
        State("improvements-rcv-loss", "value"),
        State("contents-rcv-loss", "value"),
    ],
    prevent_initial_call=True
)
def save_claim(
    n_clicks, cid,
    claim_number, policyholder, loss_address, date_of_loss, insurer, adjuster_name,
    policy_number, claim_type, contact_info_adjuster, contact_info_insured,
    coverage_a, coverage_a_deductible, coverage_a_reserve, coverage_a_advance,
    coverage_b, coverage_b_deductible, coverage_b_reserve, coverage_b_advance,
    current_claim_status_par,
    claim_assigned_date, claim_contact_date, claim_inspection_date,
    preliminary_report_par, insured_communication_paragraph, claim_reserve_paragraph,
    insured_concern_paragraph, adjuster_response_paragraph, supporting_doc_paragraph,
    next_steps_paragraph, final_report_paragraph, claim_summary_par, review_status,
    dwelling_unit_rcv_loss, detached_garage_rcv_loss,
    improvements_rcv_loss, contents_rcv_loss
):
    """
    When the user clicks 'Save Changes', update the database with the new values.
    """
    if not cid:
        return "No CID provided, cannot save."

    # Attempt to convert cid to int
    try:
        claim_id = int(cid)
    except ValueError:
        return "Invalid CID. Could not convert to int."

    conn = get_db_connection()
    cursor = conn.cursor()
    try:
        sql = """
        UPDATE claims
        SET
            claim_number = %s,
            Policyholder = %s,
            Loss_Address = %s,
            Date_Of_Loss = %s,
            Insurer = %s,
            Adjuster_Name = %s,
            Policy_Number = %s,
            claim_type = %s,
            Adjuster_Contact_Info = %s,
            Insured_Contact_Info = %s,
            coverage_building = %s,
            Coverage_A_Deductible = %s,
            Coverage_A_Reserve = %s,
            Coverage_A_Advance = %s,
            coverage_contents = %s,
            Coverage_B_Deductible = %s,
            Coverage_B_Reserve = %s,
            Coverage_B_Advance = %s,
            Current_Claim_Status_Par = %s,
            Claim_Assigned_Date = %s,
            Claim_Contact_Date = %s,
            Claim_Inspection_Date = %s,
            Preliminary_Report_Par = %s,
            Insured_Communication_Paragraph = %s,
            Claim_Reserve_Paragraph = %s,
            Insured_Concern_Paragraph = %s,
            Adjuster_Response_Paragraph = %s,
            Supporting_Doc_Paragraph = %s,
            Next_Steps_Paragraph = %s,
            Final_Report_Paragraph = %s,
            Claim_Summary_Par = %s,
            DwellingUnit_Insured_Damage_RCV_Loss = %s,
            DetachedGarage_Insured_Damage_RCV_Loss = %s,
            Improvements_Insured_Damage_RCV_Loss = %s,
            Contents_Insured_Damage_RCV_Loss = %s,
            Review_Status = %s
        WHERE id = %s
        """

        cursor.execute(sql, (
            # Existing fields
            claim_number,
            policyholder,
            loss_address,
            date_of_loss,
            insurer,
            adjuster_name,
            policy_number,
            claim_type,
            contact_info_adjuster,
            contact_info_insured,
            coverage_a,
            coverage_a_deductible,
            coverage_a_reserve,
            coverage_a_advance,
            coverage_b,
            coverage_b_deductible,
            coverage_b_reserve,
            coverage_b_advance,
            current_claim_status_par,

            # New fields
            claim_assigned_date,
            claim_contact_date,
            claim_inspection_date,
            preliminary_report_par,
            insured_communication_paragraph,
            claim_reserve_paragraph,
            insured_concern_paragraph,
            adjuster_response_paragraph,
            supporting_doc_paragraph,
            next_steps_paragraph,
            final_report_paragraph,
            claim_summary_par,
            dwelling_unit_rcv_loss,
            detached_garage_rcv_loss,
            improvements_rcv_loss,
            contents_rcv_loss,
            review_status,

            # WHERE
            claim_id
        ))
        conn.commit()
        msg = "Changes saved successfully!"
        print("Changes saved successfully!")
    except Exception as e:
        conn.rollback()
        msg = f"Error saving changes: {e}"
    finally:
        cursor.close()
        conn.close()

    return msg

def paragraph_replace_text(paragraph, regex, replace_str):
    """
    Replaces matches for 'regex' with 'replace_str' across runs in a paragraph.
    """
    while True:
        text = paragraph.text
        match = regex.search(text)
        if not match:
            break

        runs = iter(paragraph.runs)
        start, end = match.start(), match.end()

        # Advance past runs that do not contain the start of the match
        for run in runs:
            run_len = len(run.text)
            if start < run_len:
                break
            start -= run_len
            end -= run_len

        run_text = run.text
        run_len = len(run_text)
        # Replace the part of the match in this run with the replacement text
        run.text = f"{run_text[:start]}{replace_str}{run_text[end:]}"
        end -= run_len

        # Remove the remainder of the matched text from subsequent runs
        for run in runs:
            if end <= 0:
                break
            run_text = run.text
            run_len = len(run_text)
            run.text = run_text[end:]
            end -= run_len


def replace_in_paragraphs(doc, replacements):
    """
    Iterates over all paragraphs and tables in the doc, performing regex replacements.
    """
    # Replace in normal paragraphs
    for paragraph in doc.paragraphs:
        for pattern, replace_str in replacements.items():
            regex = re.compile(pattern)
            paragraph_replace_text(paragraph, regex, replace_str)

    # Replace in tables as well
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for pattern, replace_str in replacements.items():
                        regex = re.compile(pattern)
                        paragraph_replace_text(paragraph, regex, replace_str)

@callback(
    Output("download-docx", "data"),
    Output("download-notification", "children"),
    Output("download-notification", "color"),
    Input("download-docx-button", "n_clicks"),
    State("cid-store", "data"),
    prevent_initial_call=True,
)
def download_docx(n_clicks, row_id):
    if not n_clicks:
        return dash.no_update, dash.no_update, dash.no_update

    # Fetch row from DB
    conn = get_db_connection()
    cursor = conn.cursor(pymysql.cursors.DictCursor)
    
    cursor.execute("SELECT * FROM claims WHERE id = %s", (row_id,))
    row = cursor.fetchone()
    cursor.close()
    conn.close()

    if not row:
        return dash.no_update, "Claim not found!", "red"

    # Use claim_number for filename instead of row_id
    filename = f"Claim_{row['claim_number']}_Report.docx"
    
    # 2) Build a replacements dict for naive placeholders like {{Policyholder}}
    # In your Word template, you'd have placeholders literally like "{{Policyholder}}"
    replacements = {
        r"\{\{Policyholder\}\}": row.get("Policyholder", ""),
        r"\{\{DateOfLoss\}\}": row.get("Date_Of_Loss", ""),
        r"\{\{Insurer_Name\}\}": row.get("Insurer", ""),
        r"\{\{Coverage_A_Advance\}\}": row.get("Coverage_A_Advance", ""),
        r"\{\{Coverage_A_Reserve\}\}": row.get("Coverage_A_Reserve", ""),
        r"\{\{Coverage_A_Deductible\}\}": row.get("Coverage_A_Deductible", ""),
        r"\{\{coverage_building\}\}": row.get("coverage_building", ""),
        r"\{\{claim_type\}\}": row.get("claim_type", ""),
        r"\{\{Insured_Contact_Info\}\}": row.get("Insured_Contact_Info", ""),
        r"\{\{Adjuster_Contact_Info\}\}": row.get("Adjuster_Contact_Info", ""),
        r"\{\{Current_Claim_Status_Par\}\}": row.get("Current_Claim_Status_Par", ""),
        r"\{\{Claim_Assigned_Date\}\}": row.get("Claim_Assigned_Date", ""),
        r"\{\{Claim_Contact_Date\}\}": row.get("Claim_Contact_Date", ""),
        r"\{\{Claim_Inspection_Date\}\}": row.get("Claim_Inspection_Date", ""),
        r"\{\{Preliminary_Report_Par\}\}": row.get("Preliminary_Report_Par", ""),
        r"\{\{Insured_Communication_Paragraph\}\}": row.get("Insured_Communication_Paragraph", ""),
        r"\{\{Claim_Reserve_Paragraph\}\}": row.get("Claim_Reserve_Paragraph", ""),
        r"\{\{Insured_Concern_Paragraph\}\}": row.get("Insured_Concern_Paragraph", ""),
        r"\{\{Next_Steps_Paragraph\}\}": row.get("Next_Steps_Paragraph", ""),
        r"\{\{Final_Report_Paragraph\}\}": row.get("Final_Report_Paragraph", ""),
        r"\{\{Claim_Summary_Par\}\}": row.get("Claim_Summary_Par", ""),
        r"\{\{Supporting_Doc_Paragraph\}\}": row.get("Supporting_Doc_Paragraph", ""),
        r"\{\{Adjuster_Response_Paragraph\}\}": row.get("Adjuster_Response_Paragraph", ""),
        r"\{\{coverage_contents\}\}": row.get("coverage_contents", ""),
        r"\{\{Coverage_B_Deductible\}\}": row.get("Coverage_B_Deductible", ""),
        r"\{\{Coverage_B_Reserve\}\}": row.get("Coverage_B_Reserve", ""),
        r"\{\{Coverage_B_Advance\}\}": row.get("Coverage_B_Advance", ""),
        r"\{\{Adjuster_Name\}\}": row.get("Adjuster_Name", ""),
        r"\{\{Policy_Number\}\}": row.get("Policy_Number", ""),
        r"\{\{claim_number\}\}": row.get("claim_number", ""),
        r"\{\{loss_address\}\}": row.get("Loss_Address", ""),
        r"\{\{DwellingUnit_Insured_Damage_RCV_Loss\}\}": row.get("DwellingUnit_Insured_Damage_RCV_Loss", ""),
        r"\{\{DetachedGarage_Insured_Damage_RCV_Loss\}\}": row.get("DetachedGarage_Insured_Damage_RCV_Loss", ""),
        r"\{\{Improvements_Insured_Damage_RCV_Loss\}\}": row.get("Improvements_Insured_Damage_RCV_Loss", ""),
        r"\{\{Contents_Insured_Damage_RCV_Loss\}\}": row.get("Contents_Insured_Damage_RCV_Loss", ""),
        


        
        # ... add as many placeholders as you used in your .docx
        # e.g. r"\{\{Insurer\}\}": row.get("Insurer", "")
    }

    # 3) Load the Word .docx template from disk
    # Make sure you actually have this file, and your placeholders in the doc are e.g. "{{Policyholder}}"
    from docx import Document
    template_path = "/opt/PrelimSite/template.docx"
    doc = Document(template_path)

    # 4) Perform the naive placeholder replacements
    replace_in_paragraphs(doc, replacements)

    # (Optional) Sleep for debugging or demonstration
    time.sleep(2)

    # 5) Save to in-memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # 6) Return a dcc.Download object + success message
    return dcc.send_bytes(buffer.getvalue(), filename), "Report downloaded successfully.", "green"

@callback(
    Output("view-binder-link", "href"),
    Input("cid-store", "data"),
    prevent_initial_call=False
)
def update_binder_link(cid):
    print("\n" + "=" * 50)
    print(f"Callback triggered with cid: {cid}")
    
    if not cid:
        print("No CID provided")
        return "#"
    
    try:
        conn = get_db_connection()
        cursor = conn.cursor(pymysql.cursors.DictCursor)
        
        print(f"Querying database for claim_number: {cid}")
        cursor.execute("SELECT binder_spaces_link FROM claims WHERE id = %s", (cid,))
        result = cursor.fetchone()
        
        print(f"Database result: {result}")
        
        if not result or not result['binder_spaces_link']:
            print("No binder link found in database")
            return "#"
            
        spaces_link = result['binder_spaces_link']
        print(f"Found spaces link: {spaces_link}")
        
        # Parse the Spaces URL
        parsed_url = urlparse(spaces_link)
        bucket_name = 'prelim-program-file-storage'
        key = parsed_url.path.lstrip('/')
        
        print(f"Parsed bucket: {bucket_name}, key: {key}")
        
        session = boto3.session.Session()
        client = session.client('s3',
            region_name='nyc3',
            endpoint_url='https://nyc3.digitaloceanspaces.com',
            aws_access_key_id=os.getenv('SPACES_KEY'),
            aws_secret_access_key=os.getenv('SPACES_SECRET')
        )
        
        # Generate presigned URL with ResponseContentDisposition
        url = client.generate_presigned_url(
            'get_object',
            Params={
                'Bucket': bucket_name,
                'Key': key,
                'ResponseContentType': 'application/pdf',
                'ResponseContentDisposition': 'inline'  # This tells the browser to display the PDF
            },
            ExpiresIn=3600
        )
        
        print(f"Generated presigned URL: {url}")
        return url
        
    except Exception as e:
        print(f"Error in callback: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return "#"
    finally:
        if 'cursor' in locals():
            cursor.close()
        if 'conn' in locals():
            conn.close()