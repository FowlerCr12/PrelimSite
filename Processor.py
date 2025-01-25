import os
import re
import json
import zipfile
import openai
import mysql.connector
from flask import Flask, request, render_template, send_file
from docx import Document
import requests
import boto3
from io import BytesIO
from PyPDF2 import PdfReader, PdfWriter
from datetime import datetime

# Google Cloud imports
from google.cloud import documentai, storage
from google.api_core.client_options import ClientOptions

##############################################################################
#  FLASK APP SETUP
##############################################################################
app = Flask(__name__)

##############################################################################
#  CONFIG: OPENAI
##############################################################################
openai.api_key = os.getenv("OPENAI_API_KEY")

##############################################################################
#  CONFIG: GOOGLE CLOUD
##############################################################################
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "/opt/PrelimSite/service-account-key.json"  # <--- ADJUST
PROJECT_ID = "elevation-project-384914"
LOCATION = "us"  # e.g. "us" or "eu"
PROCESSOR_ID = "99a4ab70462ed46e"
GCS_BUCKET_NAME = "nfip_binder_uploads"
GCS_UPLOAD_PREFIX = "documentai/uploads/"
GCS_OUTPUT_PREFIX = "documentai/outputs/"

##############################################################################
#  CONFIG: LOCAL FOLDERS
##############################################################################
UPLOAD_FOLDER = "uploads"
JSON_OUTPUT_FOLDER = "jsonoutput"
REPORTS_FOLDER = "reports"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(JSON_OUTPUT_FOLDER, exist_ok=True)
os.makedirs(REPORTS_FOLDER, exist_ok=True)

##############################################################################
#  CONFIG: DIGITALOCEAN SPACES
##############################################################################
SPACES_ACCESS_KEY_ID = "DO801AJJWKUNKBVUBD2Q"
SPACES_SECRET_ACCESS_KEY = "VGU+W7A3zj/mVgi4AkgoO9elOIUdToeIu7hV5jBHAAc"
SPACES_ENDPOINT = "nyc3.digitaloceanspaces.com"
SPACES_BUCKET = "prelim-program-file-storage"
SPACES_REGION = "nyc3"  # adjust if needed

def get_s3_client():
    session = boto3.session.Session()
    return session.client(
        "s3",
        region_name=SPACES_REGION,
        endpoint_url=f"https://{SPACES_ENDPOINT}",
        aws_access_key_id=SPACES_ACCESS_KEY_ID,
        aws_secret_access_key=SPACES_SECRET_ACCESS_KEY
    )

def download_from_spaces(spaces_key, local_path):
    """
    Download an object from DO Spaces at 'spaces_key' to 'local_path'.
    """
    s3 = get_s3_client()
    try:
        s3.download_file(SPACES_BUCKET, spaces_key, local_path)
        print(f"[DEBUG] Downloaded from Spaces: {spaces_key} -> {local_path}")
        return True
    except Exception as e:
        print(f"[ERROR] download_from_spaces: {e}")
        return False

def upload_to_spaces(local_file_path, spaces_key):
    """
    Upload local_file_path to DO Spaces at spaces_key with ACL=public-read.
    Return the public URL if success, else None.
    """
    s3 = get_s3_client()
    try:
        s3.upload_file(local_file_path, SPACES_BUCKET, spaces_key, ExtraArgs={"ACL": "public-read"})
        url = f"https://{SPACES_BUCKET}.{SPACES_ENDPOINT}/{spaces_key}"
        print(f"[DEBUG] Uploaded to Spaces: {local_file_path} -> {url}")
        return url
    except Exception as e:
        print(f"[ERROR] upload_to_spaces: {e}")
        return None

##############################################################################
#  CONFIG: MySQL
##############################################################################
DB_HOST = "prelim-program-database-do-user-18860734-0.f.db.ondigitalocean.com"
DB_PORT = 25060  # typical DO MySQL port
DB_USER = "doadmin"
DB_PASSWORD = "AVNS_M6ypXrCa1JjgD65mbBX"
DB_NAME = "defaultdb"

def create_claims_table_if_not_exists():
    conn = mysql.connector.connect(
        host=DB_HOST,
        port=DB_PORT,
        user=DB_USER,
        password=DB_PASSWORD,
        database=DB_NAME
    )
    cursor = conn.cursor()
    create_table_sql = """
    CREATE TABLE IF NOT EXISTS claims (
      id INT AUTO_INCREMENT PRIMARY KEY,
      claim_number VARCHAR(255) NOT NULL UNIQUE,
      extracted_json JSON NOT NULL,

      Policyholder VARCHAR(255),
      Loss_Address VARCHAR(255),
      Date_Of_Loss VARCHAR(255),
      Insurer VARCHAR(255),
      Adjuster_Name VARCHAR(255),
      Policy_Number VARCHAR(255),
      Claim_Type VARCHAR(255),
      Insured_Contact_Info VARCHAR(255),
      Adjuster_Contact_Info VARCHAR(255),
      coverage_building VARCHAR(255),
      Coverage_A_Deductible VARCHAR(255),
      Coverage_A_Reserve VARCHAR(255),
      Coverage_A_Advance VARCHAR(255),
      coverage_contents VARCHAR(255),
      Coverage_B_Deductible VARCHAR(255),
      Coverage_B_Reserve VARCHAR(255),
      Coverage_B_Advance VARCHAR(255),
      Current_Claim_Status_Par VARCHAR(1000),
      Claim_Assigned_Date VARCHAR(255),
      Claim_Contact_Date VARCHAR(255),
      Claim_Inspection_Date VARCHAR(255),
      Preliminary_Report_Par VARCHAR(1000),
      Insured_Communication_Paragraph VARCHAR(1000),
      Claim_Reserve_Paragraph VARCHAR(1000),
      Insured_Concern_Paragraph VARCHAR(1000),
      Adjuster_Response_Paragraph VARCHAR(1000),
      Supporting_Doc_Paragraph VARCHAR(1000),
      Next_Steps_Paragraph VARCHAR(1000),
      Final_Report_Paragraph VARCHAR(1000),
      Claim_Summary_Par VARCHAR(1000),

      -- Additional column to store final docx link:
      report_spaces_link VARCHAR(255),

      created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
      updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
    );
    """
    cursor.execute(create_table_sql)
    cursor.close()
    conn.close()

def reformat_mdy_to_ymd(date_str):
    """
    Converts date from MM/DD/YYYY -> YYYY-MM-DD.
    If the date is empty or doesn't match MM/DD/YYYY, returns it unchanged.
    """
    if not date_str:
        return date_str
    try:
        return datetime.strptime(date_str, "%m/%d/%Y").strftime("%Y-%m-%d")
    except ValueError:
        return date_str

import json
import mysql.connector
from datetime import datetime

def reformat_mdy_to_ymd(date_str):
    """
    Converts a date from MM/DD/YYYY to YYYY-MM-DD.
    If the input is empty or doesn't match MM/DD/YYYY, returns it unchanged.
    """
    if not date_str:
        return date_str
    try:
        return datetime.strptime(date_str, "%m/%d/%Y").strftime("%Y-%m-%d")
    except ValueError:
        # If it's not in MM/DD/YYYY format (or invalid), return as-is
        return date_str

def store_claim_in_mysql(replacements, claim_number):
    """
    Insert or Update a row in 'claims' table with data from 'replacements' dict and 'claim_number'.
    Safely re-formats date fields from MM/DD/YYYY to YYYY-MM-DD to avoid MySQL date errors.
    """
    conn = mysql.connector.connect(
        host=DB_HOST,
        port=DB_PORT,
        user=DB_USER,
        password=DB_PASSWORD,
        database=DB_NAME
    )
    cursor = conn.cursor()

    insert_sql = """
    INSERT INTO claims (
        claim_number, extracted_json,
        Policyholder, Loss_Address, Date_Of_Loss, Insurer, Adjuster_Name, Policy_Number,
        Claim_Type, Insured_Contact_Info, Adjuster_Contact_Info,
        coverage_building, Coverage_A_Deductible, Coverage_A_Reserve, Coverage_A_Advance,
        coverage_contents, Coverage_B_Deductible, Coverage_B_Reserve, Coverage_B_Advance,
        Current_Claim_Status_Par, Claim_Assigned_Date, Claim_Contact_Date, Claim_Inspection_Date,
        Preliminary_Report_Par, Insured_Communication_Paragraph, Claim_Reserve_Paragraph,
        Insured_Concern_Paragraph, Adjuster_Response_Paragraph, Supporting_Doc_Paragraph,
        Next_Steps_Paragraph, Final_Report_Paragraph, Claim_Summary_Par
    )
    VALUES (
        %s, %s,
        %s, %s, %s, %s, %s, %s,
        %s, %s, %s,
        %s, %s, %s, %s,
        %s, %s, %s, %s,
        %s, %s, %s, %s,
        %s, %s, %s,
        %s, %s, %s,
        %s, %s, %s
    )
    ON DUPLICATE KEY UPDATE
      extracted_json = VALUES(extracted_json),
      Policyholder = VALUES(Policyholder),
      Loss_Address = VALUES(Loss_Address),
      Date_Of_Loss = VALUES(Date_Of_Loss),
      Insurer = VALUES(Insurer),
      Adjuster_Name = VALUES(Adjuster_Name),
      Policy_Number = VALUES(Policy_Number),
      Claim_Type = VALUES(Claim_Type),
      Insured_Contact_Info = VALUES(Insured_Contact_Info),
      Adjuster_Contact_Info = VALUES(Adjuster_Contact_Info),
      coverage_building = VALUES(coverage_building),
      Coverage_A_Deductible = VALUES(Coverage_A_Deductible),
      Coverage_A_Reserve = VALUES(Coverage_A_Reserve),
      Coverage_A_Advance = VALUES(Coverage_A_Advance),
      coverage_contents = VALUES(coverage_contents),
      Coverage_B_Deductible = VALUES(Coverage_B_Deductible),
      Coverage_B_Reserve = VALUES(Coverage_B_Reserve),
      Coverage_B_Advance = VALUES(Coverage_B_Advance),
      Current_Claim_Status_Par = VALUES(Current_Claim_Status_Par),
      Claim_Assigned_Date = VALUES(Claim_Assigned_Date),
      Claim_Contact_Date = VALUES(Claim_Contact_Date),
      Claim_Inspection_Date = VALUES(Claim_Inspection_Date),
      Preliminary_Report_Par = VALUES(Preliminary_Report_Par),
      Insured_Communication_Paragraph = VALUES(Insured_Communication_Paragraph),
      Claim_Reserve_Paragraph = VALUES(Claim_Reserve_Paragraph),
      Insured_Concern_Paragraph = VALUES(Insured_Concern_Paragraph),
      Adjuster_Response_Paragraph = VALUES(Adjuster_Response_Paragraph),
      Supporting_Doc_Paragraph = VALUES(Supporting_Doc_Paragraph),
      Next_Steps_Paragraph = VALUES(Next_Steps_Paragraph),
      Final_Report_Paragraph = VALUES(Final_Report_Paragraph),
      Claim_Summary_Par = VALUES(Claim_Summary_Par)
    """

    # Convert the replacements dict to JSON for storage
    extracted_json_str = json.dumps(replacements)

    # 1) Extract fields from replacements
    Policyholder = replacements.get("Policy_Holder", "")
    Loss_Address = replacements.get("Property_Address", "")
    raw_Date_Of_Loss = replacements.get("Date_Of_Loss", "")
    Insurer = replacements.get("Insurer_Name", "")
    Adjuster_Name = replacements.get("Adjuster_Name", "")
    Policy_Number = replacements.get("Policy_Number", "")
    claim_type = replacements.get("claim_type", replacements.get("Claim_Type", ""))
    Insured_Contact_Info = replacements.get("Policyholder_Contact_Info", "")
    Adjuster_Contact_Info = replacements.get("Adjuster_Contact_Info", "")
    coverage_building = replacements.get("Coverage-A_Building_Coverage", "")
    Coverage_A_Deductible = replacements.get("Coverage_A_Deductible", "")
    Coverage_A_Reserve = replacements.get("Coverage_A_Reserve", "")
    Coverage_A_Advance = replacements.get("Coverage_A_Advance", "")
    coverage_contents = replacements.get("Coverage-B_Contents_Coverage", "")
    Coverage_B_Deductible = replacements.get("Coverage_B_Deductible", "")
    Coverage_B_Reserve = replacements.get("Coverage_B_Reserve", "")
    Coverage_B_Advance = replacements.get("Coverage_B_Advance", "")
    Current_Claim_Status_Par = replacements.get("Claim_Status_Writeup", "")
    raw_Claim_Assigned_Date = replacements.get("Date_Assigned", "")
    raw_Claim_Contact_Date = replacements.get("Date_Contacted", "")
    raw_Claim_Inspection_Date = replacements.get("Date_Inspected", "")
    Preliminary_Report_Par = replacements.get("Preliminary_Report_Notes", "")
    Insured_Communication_Paragraph = replacements.get("Communication_With_Insured", "")
    Claim_Reserve_Paragraph = replacements.get("Claim_Reserve_Notes", "")
    Insured_Concern_Paragraph = replacements.get("Insured_Concerns", "")
    Adjuster_Response_Paragraph = replacements.get("Adj_Response_And_Comm_With_Insured", "")
    Supporting_Doc_Paragraph = replacements.get("Notes_On_Supporting_Documents", "")
    Next_Steps_Paragraph = replacements.get("Next_Claim_Steps", "")
    Final_Report_Paragraph = replacements.get("Final_Report_Summary", "")
    Claim_Summary_Par = replacements.get("Basic_Claim_Summary", "")

    # 2) Reformat dates (from MM/DD/YYYY -> YYYY-MM-DD) before inserting
    Date_Of_Loss = reformat_mdy_to_ymd(raw_Date_Of_Loss)
    Claim_Assigned_Date = reformat_mdy_to_ymd(raw_Claim_Assigned_Date)
    Claim_Contact_Date = reformat_mdy_to_ymd(raw_Claim_Contact_Date)
    Claim_Inspection_Date = reformat_mdy_to_ymd(raw_Claim_Inspection_Date)

    # 3) Pack the parameters for the INSERT statement
    data_tuple = (
        claim_number,
        extracted_json_str,
        Policyholder,
        Loss_Address,
        Date_Of_Loss,
        Insurer,
        Adjuster_Name,
        Policy_Number,
        claim_type,
        Insured_Contact_Info,
        Adjuster_Contact_Info,
        coverage_building,
        Coverage_A_Deductible,
        Coverage_A_Reserve,
        Coverage_A_Advance,
        coverage_contents,
        Coverage_B_Deductible,
        Coverage_B_Reserve,
        Coverage_B_Advance,
        Current_Claim_Status_Par,
        Claim_Assigned_Date,
        Claim_Contact_Date,
        Claim_Inspection_Date,
        Preliminary_Report_Par,
        Insured_Communication_Paragraph,
        Claim_Reserve_Paragraph,
        Insured_Concern_Paragraph,
        Adjuster_Response_Paragraph,
        Supporting_Doc_Paragraph,
        Next_Steps_Paragraph,
        Final_Report_Paragraph,
        Claim_Summary_Par
    )

    # 4) Execute INSERT / UPDATE
    cursor.execute(insert_sql, data_tuple)
    conn.commit()

    # 5) Close
    cursor.close()
    conn.close()


##############################################################################
# GOOGLE CLOUD DOCUMENT AI & UTILS
##############################################################################
def upload_to_gcs(file_path, bucket_name, prefix):
    storage_client = storage.Client()
    bucket = storage_client.bucket(bucket_name)
    blob_name = os.path.join(prefix, os.path.basename(file_path))
    blob = bucket.blob(blob_name)
    blob.upload_from_filename(file_path)
    return f"gs://{bucket_name}/{blob_name}"

def process_with_document_ai(input_gcs_uri, output_gcs_uri):
    try:
        client = documentai.DocumentProcessorServiceClient(
            client_options=ClientOptions(api_endpoint=f"{LOCATION}-documentai.googleapis.com")
        )
        processor_name = client.processor_path(PROJECT_ID, LOCATION, PROCESSOR_ID)

        input_config = documentai.BatchDocumentsInputConfig(
            gcs_prefix=documentai.GcsPrefix(gcs_uri_prefix=input_gcs_uri)
        )
        output_config = documentai.DocumentOutputConfig(
            gcs_output_config=documentai.DocumentOutputConfig.GcsOutputConfig(
                gcs_uri=output_gcs_uri
            )
        )

        request = documentai.BatchProcessRequest(
            name=processor_name,
            input_documents=input_config,
            document_output_config=output_config,
        )
        operation = client.batch_process_documents(request=request)
        operation.result(timeout=600)
        return True
    except Exception as e:
        print(f"Error processing with Document AI: {e}")
        return False

def download_from_gcs(gcs_prefix, local_folder):
    storage_client = storage.Client()
    short_prefix = gcs_prefix.replace(f"gs://{GCS_BUCKET_NAME}/", "")
    bucket = storage_client.bucket(GCS_BUCKET_NAME)
    blobs = bucket.list_blobs(prefix=short_prefix)

    local_files = []
    os.makedirs(local_folder, exist_ok=True)
    for blob in blobs:
        if blob.name.endswith(".json"):
            local_path = os.path.join(local_folder, os.path.basename(blob.name))
            blob.download_to_filename(local_path)
            local_files.append(local_path)
    return local_files

def read_text_file(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

def call_openai_for_json(input_text, custom_prompt):
    try:
        response = openai.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": custom_prompt},
                {"role": "user", "content": input_text},
            ],
            max_tokens=3500,
            temperature=0.1
        )
        ai_raw_text = response.choices[0].message.content.strip()
        data = json.loads(ai_raw_text)
        return data
    except json.JSONDecodeError as e:
        print("Error: The response from OpenAI was not valid JSON.")
        print("AI raw text:", ai_raw_text)
        raise e
    except Exception as e:
        print(f"Error calling OpenAI: {e}")
        raise e

##############################################################################
#  DOCX REPLACEMENT
##############################################################################
def paragraph_replace_text(paragraph, regex, replace_str):
    """
    Replaces matches for 'regex' with 'replace_str' across runs in a paragraph.
    """
    while True:
        text = paragraph.text
        match = regex.search(text)
        if not match:
            break

        runs = iter(paragraph.runs)
        start, end = match.start(), match.end()

        # Advance past runs that do not contain the start of the match
        for run in runs:
            run_len = len(run.text)
            if start < run_len:
                break
            start -= run_len
            end -= run_len

        run_text = run.text
        run_len = len(run_text)
        # Replace the part of the match in this run with the replacement text
        run.text = f"{run_text[:start]}{replace_str}{run_text[end:]}"
        end -= run_len

        # Remove the remainder of the matched text from subsequent runs
        for run in runs:
            if end <= 0:
                break
            run_text = run.text
            run_len = len(run_text)
            run.text = run_text[end:]
            end -= run_len

    return paragraph


def replace_in_paragraphs(doc, replacements):
    """
    Iterates over all paragraphs and tables in the doc, performing regex replacements.
    """
    for paragraph in doc.paragraphs:
        for pattern, replace_str in replacements.items():
            regex = re.compile(pattern)
            paragraph_replace_text(paragraph, regex, replace_str)

    # If placeholders might appear in tables, handle them too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for pattern, replace_str in replacements.items():
                        regex = re.compile(pattern)
                        paragraph_replace_text(paragraph, regex, replace_str)

##############################################################################
#  PROCESS CLAIM PAIR
##############################################################################
def process_claim_pair(pdf_path, txt_path, claim_number):
    """
    1) Upload PDF to GCS
    2) Document AI
    3) Download JSON
    4) read .txt => OpenAI
    5) Merge data
    6) store_claim_in_mysql
    7) docx replace
    8) produce .docx => return path
    """
    base_name = claim_number  # or os.path.splitext(...) but we have direct claim_number

    # 1) Upload PDF -> GCS
    input_gcs_uri = upload_to_gcs(pdf_path, GCS_BUCKET_NAME, GCS_UPLOAD_PREFIX)
    output_gcs_uri = f"gs://{GCS_BUCKET_NAME}/{GCS_OUTPUT_PREFIX}{base_name}/"
    print("[DEBUG] Uploaded PDF to GCS:", input_gcs_uri)
    print("[DEBUG] Document AI output URI:", output_gcs_uri)

    # 2) Document AI
    success = process_with_document_ai(input_gcs_uri, output_gcs_uri)
    if not success:
        raise Exception("Document AI processing failed.")

    # 3) Download JSON
    output_files = download_from_gcs(output_gcs_uri, JSON_OUTPUT_FOLDER)
    if not output_files:
        raise Exception("No JSON outputs from Document AI.")
    with open(output_files[0], "r", encoding="utf-8") as f:
        docai_data = json.load(f)

    # 4) read .txt => OpenAI
    text_content = read_text_file(txt_path)
    if not text_content.strip():
        raise Exception("The .txt file is empty.")
    custom_prompt = """
    You are DaveGPT, the world's best flood insurance adjuster, that only outputs valid JSON, with no Markdown formatting or code fences, in the following format:
    {
      "entities": [
        { "type": "...", "mentionText": "..." },
        ...
      ]
    }
    Do NOT include triple backticks or any extra text. Please take into account the timeline of the claim notes when crafting your response in order to ensure that you return the most accurate and up to date information. The information you provide must be 100% accurate, due to the important nature of insurance. Also note, the acronym FA means Field Adjuster, who is the main adjuster on the claim. Additionally, avoid talking about topics of personal matters of the adjuster. If there are issues regarding the claim because of issues with the adjuster, simply note that the claim is delayed, etc, but do not go too into depth. If the notes of the claim contain information regarding the use of a public adjuster please ensure you not this at some point. 
    Provide information based on claim notes. The information will be separated into separate sections, with no more than 3-4 sentences for each section, and only containing the most relevant information. Do not include sources, but ensure that the information you provide is accurate. Only provide the information requested, and nothing more in your responses. Under no circumstances should you include dollar amounts or figures. Your job is solely to summarize the notes and important information regarding the claim. The entities you create should be of the types: "Claim_Status_Writeup", "Preliminary_Report_Notes", "Communication_With_Insured", "Claim_Reserve_Notes", "Insured_Concerns", "Adj_Response_And_Comm_With_Insured", "Notes_On_Supporting_Documents", "Next_Claim_Steps", "Basic_Claim_Summary", and "Final_Report_Summary".
    For "Claim_Status_Writeup" include a very brief 1-3 sentence summary on where at in the claims process the current claim is at.
    For "Preliminary_Report_Notes", provide a short summary on the status of the preliminary report as well as what is currently going on/has occurred regarding the preliminary report.
    For "Communication_With_Insured", provide a short summary on how the adjuster and insured have been communicating, and what they have talked about regarding the overall claim.
    For "Claim_Reserve_Notes", provide a short summary on the status of the claim reserves, how they have been adjusted (if at all), and the next steps in setting the reserves or finalizing them.
    For "Insured_Concerns", provide a short summary on any concerns the insured has mentioned, as well as how the adjuster has tried to alleviate these concerns.
    For "Adj_Response_And_Comm_With_Insured" provide a short summary on how the adjuster has communicated with the insured to ensure the claim is completed satisfactorily. If there have been issues with communication from either side (adjuster & insured), be sure to note this.
    For "Notes_On_Supporting_Documents" provide a short summary on the supporting claim documents the insured has provided, whether that be pictures, videos, lease agreements, mortgage agreements, or anything relevant to the claim. If a public adjuster (PA) has been used, note this here. 
    For "Next_Claim_Steps" provide a short summary detailing the next steps the adjuster is taking/working through in order to ensure the claim comes to a close.
    For "Basic_Claim_Summary" provide quick summary on the information that has been looked through and summarized in previous sections.
    For "Final_Report_Summary" provide a quick summary of the overall final report which is found at the end of the information supplied to you. Additionally, make sure you write information on denials if there were any in the claim. Provide a short snippet on the amount of the estimate as well. Make sure that the "Insured Damage RCV loss" is provided here for both Coverage A and Coverage B (if claim contains Coverage B). The RCV is provided in three numbers. The Insured damage RCV loss is very important. Make sure not to get it confused with any other numbers. (i.e. Property pre-loss value RCV). Note, there will not always be a final report given to you. If this is the case, simply enter the information "No Final Report at this time." For this json entry. 
    Final note: PLease ensure you are thorough and 100% accurate. It is imperative that the information you provide is 100% correct. Additionally, if you have trouble finding the final report, if there even is one, it will be marked with "===== Final Report Text (First 7 Pages) =====" at the start of the information. 
    """
    ai_data = call_openai_for_json(text_content, custom_prompt)
    print("AI Data:", ai_data)

    # 5) Merge docai_data + ai_data
    replacements = {}
    # from docai
    for e in docai_data.get("entities", []):
        t = e.get("type", "")
        mention = e.get("mentionText", "")
        if t:
            replacements[t] = mention
    # from openai
    for e in ai_data.get("entities", []):
        t = e.get("type", "")
        mention = e.get("mentionText", "")
        if t:
            replacements[t] = mention

    print("success 7")

    # Adjuster & Insurer
    adjuster_name = replacements.get("Adjuster", "NO_ADJUSTER")
    insurer_name = replacements.get("Insurer", "NO_INSURER")
    replacements["Adjuster_Name"] = adjuster_name
    replacements["Insurer_Name"] = insurer_name
    print("success 8")

    # Contact info
    policyholder_email = replacements.get("Policyholder_Email", "")
    policyholder_phone = replacements.get("Policyholder_Phone", "")
    if not policyholder_email.strip() and not policyholder_phone.strip():
        policyholder_contact_info = "!!!Not Provided in Prelim.!!!"
    else:
        policyholder_contact_info = policyholder_email if policyholder_email else policyholder_phone
    replacements["Policyholder_Contact_Info"] = policyholder_contact_info
    print("success 9")
    print("success 10")

    # Remove the original keys from docAI to avoid double replacements
    replacements.pop("Adjuster", None)
    replacements.pop("Insurer", None)
    print("success 11")

    # store in MySQL
    store_claim_in_mysql(replacements, claim_number)

    # 7) docx replace
    docx_template = "/opt/PrelimScraper/template.docx"  # <--- ADJUST path
    if not os.path.isfile(docx_template):
        raise Exception("template.docx not found.")
    doc = Document(docx_template)
    replace_in_paragraphs(doc, replacements)

    # 8) produce docx
    output_docx_path = os.path.join(REPORTS_FOLDER, f"{base_name}.docx")
    doc.save(output_docx_path)
    print("[DEBUG] Created final docx:", output_docx_path)
    return output_docx_path

##############################################################################
#  FLASK ROUTES
##############################################################################

@app.route("/")
def index():
    return """<h1>Welcome</h1><p>Use /process_db or /process to proceed.</p>"""

@app.route("/process_db", methods=["GET"])
def process_from_db():
    """
    1) Query claims table for rows with pdf/notes in Spaces,
    2) Download them,
    3) process_claim_pair(),
    4) Upload final docx to 'Reports/',
    5) update DB with report_spaces_link
    """
    create_claims_table_if_not_exists()

    conn = mysql.connector.connect(
        host=DB_HOST, port=DB_PORT, user=DB_USER, password=DB_PASSWORD, database=DB_NAME
    )
    cursor = conn.cursor(dictionary=True)

    select_sql = """
    SELECT *
    FROM claims
    WHERE binder_spaces_link IS NOT NULL
      AND binder_spaces_link != ''
      AND notes_spaces_link IS NOT NULL
      AND notes_spaces_link != ''
      AND (report_spaces_link IS NULL OR report_spaces_link = '')
    """
    cursor.execute(select_sql)
    rows = cursor.fetchall()

    processed_claims = []
    for row in rows:
        dox_path = None
        claim_number = row["claim_number"]
        pdf_spaces_link = row.get("binder_spaces_link", "")
        notes_spaces_link = row.get("notes_spaces_link", "")

        # parse object keys from URLs
        pdf_key = pdf_spaces_link.replace(f"https://{SPACES_BUCKET}.{SPACES_ENDPOINT}/", "")
        notes_key = notes_spaces_link.replace(f"https://{SPACES_BUCKET}.{SPACES_ENDPOINT}/", "")

        local_pdf = os.path.join(UPLOAD_FOLDER, f"{claim_number}.pdf")
        local_txt = os.path.join(UPLOAD_FOLDER, f"{claim_number}.txt")

        ok_pdf = download_from_spaces(pdf_key, local_pdf)
        ok_txt = download_from_spaces(notes_key, local_txt)
        if not (ok_pdf and ok_txt):
            print(f"[ERROR] Could not download PDF or TXT for claim {claim_number}")
            continue

        docx_path = None  # Initialize before try block
        try:
            docx_path = process_claim_pair(local_pdf, local_txt, claim_number)
            # upload final docx
            docx_filename = os.path.basename(docx_path)
            docx_key = f"Reports/{docx_filename}"
            docx_url = upload_to_spaces(docx_path, docx_key)

            if docx_url:
                # update DB
                p_sql = """
                    UPDATE claims
                    SET report_spaces_link = %s,
                    Review_Status = 'In Review'
                    WHERE claim_number = %s
                """
                cursor.execute(up_sql, (docx_url, claim_number))
                conn.commit()
                processed_claims.append((claim_number, docx_url))

        except Exception as e:
            print(f"[ERROR] processing claim {claim_number}: {e}")
        finally:
            # cleanup local
            if os.path.exists(local_pdf):
                os.remove(local_pdf)
            if os.path.exists(local_txt):
                os.remove(local_txt)
            if docx_path and os.path.exists(docx_path):  # Only try to remove if it exists and is not None
                os.remove(docx_path)

    cursor.close()
    conn.close()

    resp = "Processed:\n"
    for cnum, link in processed_claims:
        resp += f"Claim {cnum} => {link}\n"
    return resp, 200

@app.route("/process", methods=["POST"])
def process_upload():
    """
    Original approach: user uploads PDF+TXT, we pair, process, zip docx.
    Adjust or remove as needed.
    """
    create_claims_table_if_not_exists()

    all_files = request.files.getlist("files")
    saved_files = []
    for f in all_files:
        if f.filename:
            path = os.path.join(UPLOAD_FOLDER, f.filename)
            f.save(path)
            saved_files.append(path)

    claims_map = {}
    for path in saved_files:
        base, ext = os.path.splitext(os.path.basename(path))
        ext = ext.lower()
        if base not in claims_map:
            claims_map[base] = {}
        if ext == ".pdf":
            claims_map[base]["pdf"] = path
        elif ext == ".txt":
            claims_map[base]["txt"] = path

    processed_reports = []
    for base_name, file_dict in claims_map.items():
        pdf_path = file_dict.get("pdf")
        txt_path = file_dict.get("txt")
        if not pdf_path or not txt_path:
            print(f"Skipping {base_name}; missing PDF or TXT.")
            continue

        try:
            docx_path = process_claim_pair(pdf_path, txt_path, base_name)
            processed_reports.append(docx_path)
        except Exception as e:
            print(f"Error: {e}")

        # cleanup
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
        if os.path.exists(txt_path):
            os.remove(txt_path)

    zip_path = os.path.join(REPORTS_FOLDER, "reports.zip")
    with zipfile.ZipFile(zip_path, "w") as zipf:
        for rep in processed_reports:
            zipf.write(rep, os.path.basename(rep))
            os.remove(rep)

    return send_file(zip_path, as_attachment=True)

@app.route("/uploadpage", methods=["GET"])
def upload_page():
    """
    A simple HTML form for manual upload (optional).
    """
    return """
    <html><body>
    <h1>Upload PDF + TXT</h1>
    <form action="/process" method="post" enctype="multipart/form-data">
      <input type="file" name="files" multiple>
      <button type="submit">Upload & Process</button>
    </form>
    </body></html>
    """

##############################################################################
#  RUN FLASK
##############################################################################
if __name__ == "__main__":
    create_claims_table_if_not_exists()
    app.run(host="0.0.0.0", port=5003, debug=True)
